from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
REPORT_MD_PATH = ROOT / "report.md"
ASSETS_DIR = ROOT / "assets"
DOCX_PATH = ROOT / "消费红利策略研究报告.docx"
DOCX_FALLBACK_PATH = ROOT / "消费红利策略研究报告_v2.docx"

BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)
ACCENT_COLOR = RGBColor(0x00, 0x00, 0x00)
HEADER_FILL = "EDEDED"
ROW_FILL = "F7F7F7"
INFO_LEFT_FILL = "EEEEEE"
INFO_RIGHT_FILL = "FAFAFA"
BODY_LINE_SPACING_PT = 20
MAX_IMAGE_WIDTH_CM = 15.2
MAX_IMAGE_HEIGHT_CM = 18.8

FIGURE_CAPTIONS = {
    "fig_01_monthly_valid.png": (
        "有效样本数量变化",
        "展示每月可交易股票数量的动态扩容过程，说明本文采用的是月度动态样本池而非固定成分池。",
    ),
    "fig_02_divi_dist.png": (
        "divi 字段分布与口径验证",
        "展示股息率快照的横截面分布及右偏特征，为后续采用稳健排序和规模中性化处理提供依据。",
    ),
    "fig_03_group_nav.png": (
        "核心因子五分组净值",
        "比较各因子从高分组到低分组的净值分化，用于判断因子是否具备单调性和可交易的横截面区分度。",
    ),
    "fig_04_ic_series.png": (
        "因子 IC 序列与累计 IC",
        "展示因子预测下月收益的时间稳定性；累计 IC 持续上行代表信号方向较稳定。",
    ),
    "fig_05_correlation.png": (
        "因子相关性矩阵",
        "检验红利、规模、反转和低波因子之间的信息重叠程度，辅助判断多因子合成是否存在信号稀释。",
    ),
    "fig_06_nav.png": (
        "策略净值对比",
        "比较核心策略、增强策略和等权基准的长期净值表现，是判断策略收益贡献的主图。",
    ),
    "fig_07_excess.png": (
        "相对基准超额净值",
        "展示策略相对消费行业等权基准的超额收益累积路径，用于识别策略有效和失效的市场阶段。",
    ),
    "fig_08_drawdown.png": (
        "策略回撤对比",
        "比较策略与基准的回撤深度和修复速度，刻画红利策略的防御属性与尾部风险。",
    ),
    "fig_09_annual.png": (
        "年度收益拆分",
        "按自然年展示策略、基准和超额收益，说明策略在弱市中更容易体现防御收益，在牛市中可能跑输。",
    ),
    "fig_10_ls_nav.png": (
        "多空组合净值",
        "通过做多高得分组、做空低得分组观察纯因子收益，剥离市场方向后检验因子本身的有效性。",
    ),
    "fig_11_buffer.png": (
        "再平衡缓冲带效果",
        "比较缓冲带前后的净值与换手，检验降换手机制是否以过多收益损失为代价。",
    ),
    "fig_12_bootstrap.png": (
        "Bootstrap 显著性检验",
        "展示夏普差异的重采样分布，用于评估策略相对基准优势在小样本下的统计稳健性。",
    ),
    "fig_13_random_benchmark.png": (
        "随机组合基准分布",
        "用 1000 组随机 Top15 组合形成经验分布，检验策略收益是否明显优于随机选股结果。",
    ),
}

TABLE_CAPTIONS = [
    ("样本概况", "概括原始数据规模、时间范围和动态有效样本变化。"),
    ("divi 字段口径验证", "检验 divi 是否符合月度股息率快照而非分红事件标识。"),
    ("DY 构建路线验证", "比较直接填充 divi 与 DPS 代理值前向填充后的差异。"),
    ("候选因子库", "列示由现有字段可实现的因子及处理意见。"),
    ("单因子检验结果", "通过多空收益、IC 和 RankIC 评估因子的横截面预测能力。"),
    ("回测协议", "锁定投资域、信号时间、调仓、费用、基准和收益口径。"),
    ("策略绩效比较", "集中比较主策略、候选策略和等权基准的收益、风险与换手。"),
    ("低换手增强方案", "比较 DY 规模中性、等权反转合成和温和混合方案的收益换手平衡。"),
    ("随机组合基准", "用随机 Top15 经验分布评估策略收益是否可能来自小样本运气。"),
    ("Bootstrap 夏普差异检验", "用重采样置信区间评估策略相对基准的夏普优势。"),
    ("Fama-MacBeth 单因子回归", "从截面回归角度评估单因子风险溢价及其显著性。"),
    ("Fama-MacBeth 多因子回归", "在控制规模和反转后检验 DY 的边际独立贡献。"),
    ("子期策略表现", "比较前后两个子期中 DY 策略的收益和风险。"),
    ("调仓频率敏感性", "比较月度与季度调仓对收益、风险和换手的影响。"),
    ("TopN 敏感性", "检验持股数量变化对收益和分散化效果的影响。"),
    ("费率敏感性", "评估交易费率上升对策略绩效的侵蚀程度。"),
    ("再平衡缓冲带", "检验降低换手是否需要牺牲过多收益。"),
    ("动态 TopN", "检验信号离散度调仓规则是否优于固定持股数量。"),
    ("滚动验证", "用训练期选因子、验证期回测的方式检验因子选择稳定性。"),
]


def set_run_font(
    run,
    size_pt: float = 12,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    latin_font: str = "Times New Roman",
    east_asia_font: str = "宋体",
):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    run.font.color.rgb = color or BODY_COLOR

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def clean_inline(text: str) -> str:
    text = text.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def add_inline_runs(
    paragraph,
    text: str,
    *,
    size_pt: float = 12,
    base_bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    east_asia_font: str = "宋体",
):
    """Render **bold** and `code` markers while keeping the docx stable."""
    text = clean_inline(text)
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(
                run,
                size_pt=size_pt,
                bold=base_bold,
                italic=italic,
                color=color,
                east_asia_font=east_asia_font,
            )

        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2].replace("`", ""))
            set_run_font(
                run,
                size_pt=size_pt,
                bold=True,
                italic=italic,
                color=color,
                east_asia_font=east_asia_font,
            )
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                size_pt=size_pt - 0.5,
                color=MUTED_COLOR,
                latin_font="Consolas",
                east_asia_font="宋体",
            )
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(
            run,
            size_pt=size_pt,
            bold=base_bold,
            italic=italic,
            color=color,
            east_asia_font=east_asia_font,
        )


def set_paragraph_bottom_border(paragraph, color: str = "808080", size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 80, right: int = 80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_compact_spacing(paragraph, *, before: float = 0, after: float = 0):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)

    head = paragraph.add_run("第 ")
    set_run_font(head, size_pt=9.5, color=MUTED_COLOR)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    set_run_font(field_run, size_pt=9.5, color=MUTED_COLOR)
    field_run._r.append(begin)
    field_run._r.append(instr)
    field_run._r.append(end)

    tail = paragraph.add_run(" 页")
    set_run_font(tail, size_pt=9.5, color=MUTED_COLOR)


def add_header(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("消费红利策略研究")
    set_run_font(run, size_pt=9.5, color=MUTED_COLOR)


def configure_document(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BODY_COLOR
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.55)
    section.right_margin = Cm(2.55)

    add_header(section.header.paragraphs[0])
    add_page_number(section.footer.paragraphs[0])


def add_horizontal_rule(doc: Document, color: str = "000000", size: str = "10"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_bottom_border(p, color=color, size=size)


def add_cover_page(doc: Document):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("消费红利策略研究报告")
    set_run_font(run, size_pt=24, bold=True, color=ACCENT_COLOR, east_asia_font="黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("红利因子的选股能力与低换手增强")
    set_run_font(run, size_pt=15, color=MUTED_COLOR, east_asia_font="黑体")

    add_horizontal_rule(doc)

    items = [
        ("研究对象", "A 股消费行业月频横截面选股"),
        ("数据区间", "2013 年 8 月至 2022 年 11 月"),
        ("策略口径", "DY（规模中性）+ 反转辅助增强"),
        ("评估维度", "因子有效性、组合回测、稳健性、统计验证"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for left, right in items:
        row = table.add_row().cells
        set_cell_shading(row[0], INFO_LEFT_FILL)
        set_cell_shading(row[1], INFO_RIGHT_FILL)
        for cell in row:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(left)
        set_run_font(r0, size_pt=11, bold=True, east_asia_font="黑体")
        p1 = row[1].paragraphs[0]
        add_inline_runs(p1, right, size_pt=11)

    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("研究报告最终版")
    set_run_font(run, size_pt=12, color=MUTED_COLOR)

    doc.add_page_break()


def add_navigation_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("报告结构")
    set_run_font(run, size_pt=18, bold=True, east_asia_font="黑体")

    sections = [
        "摘要：核心结论与收益口径",
        "1. 数据与口径",
        "2. 因子有效性",
        "3. 策略回测",
        "4. 稳健性与统计验证",
        "5. 组合定位",
        "6. 数据限制与风险提示",
        "附录：敏感性测试与补充图表",
    ]
    for item in sections:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2.0)
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, item, size_pt=12.5, base_bold=True)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int):
    if level == 2 and re.match(r"^\d+\.\s", text) and not text.startswith("1. "):
        doc.add_section(WD_SECTION_START.NEW_PAGE)

    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(16)
        add_inline_runs(p, text, size_pt=20, base_bold=True, east_asia_font="黑体")
        return

    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(6 if level in (2, 3) else 3)
    if level == 2:
        add_inline_runs(p, text, size_pt=14.5, base_bold=True, east_asia_font="黑体")
        set_paragraph_bottom_border(p)
    elif level == 3:
        add_inline_runs(p, text, size_pt=12.5, base_bold=True, east_asia_font="黑体")
    else:
        add_inline_runs(p, text, size_pt=11.5, base_bold=True, color=MUTED_COLOR, east_asia_font="黑体")


def add_text_paragraph(doc: Document, text: str, *, center: bool = False, size_pt: float = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    p.paragraph_format.space_after = Pt(6)
    if not center:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_runs(p, text, size_pt=size_pt)


def add_lead_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text, size_pt=11.5, italic=True, color=MUTED_COLOR)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    p.paragraph_format.space_after = Pt(3)
    add_inline_runs(p, text, size_pt=12)


def add_numbered_item(doc: Document, number: str, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.44)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    p.paragraph_format.space_after = Pt(3)
    add_inline_runs(p, f"{number}. {text}", size_pt=12)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_compact_spacing(p, after=6)
    add_inline_runs(p, text, size_pt=10.5, base_bold=True, color=MUTED_COLOR)


def add_figure_caption(doc: Document, figure_no: int, title: str, note: str):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.keep_together = True
    title_p.paragraph_format.keep_with_next = True
    set_compact_spacing(title_p, after=1)
    add_inline_runs(title_p, f"图 {figure_no} {title}", size_pt=10.5, base_bold=True, color=ACCENT_COLOR)

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note_p.paragraph_format.left_indent = Cm(0.75)
    note_p.paragraph_format.right_indent = Cm(0.75)
    note_p.paragraph_format.keep_together = True
    set_compact_spacing(note_p, after=8)
    add_inline_runs(note_p, f"说明：{note}", size_pt=10, color=MUTED_COLOR)


def add_table_caption(doc: Document, table_no: int):
    if table_no <= len(TABLE_CAPTIONS):
        title, note = TABLE_CAPTIONS[table_no - 1]
    else:
        title, note = "补充表格", "汇总报告正文中的补充数据。"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    set_compact_spacing(p, before=4, after=4)
    add_inline_runs(p, f"表 {table_no} {title}：{note}", size_pt=10.5, base_bold=True, color=MUTED_COLOR)


def resolve_image(rel_path: str) -> Path:
    raw = rel_path.strip().replace("\\", "/")
    path = ROOT / raw
    if path.exists():
        return path
    path = ASSETS_DIR / Path(raw).name
    if path.exists():
        return path
    raise FileNotFoundError(f"Image referenced in report.md was not found: {rel_path}")


def image_display_size_cm(img_path: Path) -> tuple[float, float | None]:
    with Image.open(img_path) as img:
        width_px, height_px = img.size
    height_at_max_width = MAX_IMAGE_WIDTH_CM * height_px / width_px
    if height_at_max_width <= MAX_IMAGE_HEIGHT_CM:
        return MAX_IMAGE_WIDTH_CM, None
    width_at_max_height = MAX_IMAGE_HEIGHT_CM * width_px / height_px
    return width_at_max_height, MAX_IMAGE_HEIGHT_CM


def add_image(doc: Document, img_path: Path, figure_no: int, alt_text: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    set_compact_spacing(p, before=4, after=2)
    run = p.add_run()
    width_cm, height_cm = image_display_size_cm(img_path)
    if height_cm is None:
        run.add_picture(str(img_path), width=Cm(width_cm))
    else:
        run.add_picture(str(img_path), height=Cm(height_cm))

    title, note = FIGURE_CAPTIONS.get(
        img_path.name,
        (alt_text or img_path.stem, "展示报告正文对应分析结果。"),
    )
    add_figure_caption(doc, figure_no, title, note)


def is_table_block(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    current = lines[idx].strip()
    nxt = lines[idx + 1].strip()
    return current.startswith("|") and nxt.startswith("|") and set(nxt.replace("|", "").strip()) <= {"-", ":", " "}


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown_table(lines: list[str], idx: int):
    header = split_table_row(lines[idx])
    rows = []
    i = idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        row = split_table_row(line)
        if len(row) < len(header):
            row += [""] * (len(header) - len(row))
        rows.append(row[: len(header)])
        i += 1
    return header, rows, i


def sanitize_table_text(text: str) -> str:
    return clean_inline(text).replace("**", "").replace("`", "")


def add_table(doc: Document, header: list[str], rows: list[list[str]], table_no: int):
    add_table_caption(doc, table_no)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cols = len(header)
    font_size = 10.5 if cols <= 3 else 9.5 if cols <= 6 else 8.5
    is_info_table = [sanitize_table_text(x) for x in header] == ["项目", "内容"]

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell, top=70, bottom=70, left=60, right=60)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_compact_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(
            p,
            sanitize_table_text(text),
            size_pt=font_size,
            base_bold=True,
            east_asia_font="黑体",
        )

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cell = cells[j]
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_info_table:
                set_cell_shading(cell, INFO_LEFT_FILL if j == 0 else INFO_RIGHT_FILL)
            elif row_idx % 2 == 1:
                set_cell_shading(cell, ROW_FILL)
            p = cell.paragraphs[0]
            set_compact_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, sanitize_table_text(text), size_pt=font_size)

    spacer = doc.add_paragraph()
    set_compact_spacing(spacer, after=2)


def build_docx() -> Path:
    if not REPORT_MD_PATH.exists():
        raise FileNotFoundError(f"Cannot find {REPORT_MD_PATH}")

    lines = REPORT_MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover_page(doc)
    add_navigation_page(doc)

    paragraph_buffer: list[str] = []
    figure_no = 0
    table_no = 0

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(x.strip() for x in paragraph_buffer if x.strip())
            if text:
                add_text_paragraph(doc, text)
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped in {"---", "***", "___"}:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        if stripped.startswith("#### "):
            flush_paragraph()
            add_heading(doc, stripped[5:].strip(), 4)
            i += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_lead_line(doc, stripped[2:].strip())
            i += 1
            continue

        if is_table_block(lines, i):
            flush_paragraph()
            header, rows, next_idx = parse_markdown_table(lines, i)
            table_no += 1
            add_table(doc, header, rows, table_no)
            i = next_idx
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            alt_text = image_match.group(1).strip()
            img_path = resolve_image(image_match.group(2))
            figure_no += 1
            add_image(doc, img_path, figure_no, alt_text=alt_text)
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_bullet(doc, stripped[2:].strip())
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if number_match:
            flush_paragraph()
            add_numbered_item(doc, number_match.group(1), number_match.group(2).strip())
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        doc.save(DOCX_FALLBACK_PATH)
        return DOCX_FALLBACK_PATH


if __name__ == "__main__":
    output = build_docx()
    print(f"Generated {output.relative_to(ROOT)}")
